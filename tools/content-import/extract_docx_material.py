#!/usr/bin/env python3
"""Extract paragraphs/tables from the user-provided DOCX lesson archive.

The extractor never decides which spelling is authoritative.  It creates a
traceable text snapshot that the pack builder can cite in source_location.
"""
from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document


PRIVATE_REPLACEMENTS = {
    "Leonardo": "Bambino", "Leo": "Aram", "Sona": "Mariam",
    "Paolo": "Genitore", "Alessandra": "Genitore", "Arthur": "Davit",
    "Freya": "Ani", "Spritz": "Micio", "Steen": "Adulto",
    "Լեոնարդո": "Արամ", "Սոնա": "Մարիամ", "Պաոլո": "Ծնող",
    "Արթուր": "Դավիթ", "Ալեսանդրա": "Ծնող", "Ֆրեյա": "Անի"
}

def normalize(text: str) -> str:
    value = " ".join(text.replace("\u00a0", " ").split())
    for source, replacement in PRIVATE_REPLACEMENTS.items():
        value = value.replace(source, replacement)
    return value


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    lines: list[str] = []
    for path in sorted(args.input.glob("*.docx")):
        lines.append(f"===== {path.name} =====")
        doc = Document(path)
        for index, paragraph in enumerate(doc.paragraphs):
            text = normalize(paragraph.text)
            if text:
                lines.append(f"P{index:03d}\t{text}")
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                cells = [normalize(cell.text) for cell in row.cells]
                if any(cells):
                    lines.append(f"T{table_index}.{row_index}\t" + " | ".join(cells))
        lines.append("")
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {args.output} ({len(lines)} records)")


if __name__ == "__main__":
    main()
